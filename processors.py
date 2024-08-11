import asyncio
from PyPDF2 import PdfReader
import docx
import requests
from bs4 import BeautifulSoup
import speech_recognition as sr
from pydub import AudioSegment
from io import BytesIO
from utils import extract_text_from_image, summarize_text, generate_embedding, store_in_vector_db, search_vector_db, query_knowledge_base

async def process_document(doc, username):
    file = await doc.get_file()
    file_content = await file.download_as_bytearray()
    
    if doc.file_name.endswith('.pdf'):
        return await process_pdf(BytesIO(file_content), username)
    elif doc.file_name.endswith('.docx'):
        return await process_docx(BytesIO(file_content), username)
    else:
        return "Unsupported document type"

async def process_pdf(file_obj, username):
    reader = PdfReader(file_obj)
    if len(reader.pages) > 30:
        return "PDF is too long (over 30 pages)"
    
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    
    if not text:
        # If no text extracted, it might be a scanned PDF
        return await process_scanned_pdf(file_obj, username)
    
    summary = await summarize_text(text)
    embedding = await generate_embedding(summary)
    await store_in_vector_db(embedding, summary, username)
    return "PDF processed and stored"

async def process_scanned_pdf(file_obj, username):
    # Use OCR or GPT-4 Vision here
    text = await extract_text_from_image(file_obj)
    summary = await summarize_text(text)
    embedding = await generate_embedding(summary)
    await store_in_vector_db(embedding, summary, username)
    return "Scanned PDF processed and stored"

async def process_docx(file_obj, username):
    doc = docx.Document(file_obj)
    
    # Count total words
    total_words = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
    
    # Estimate pages (assuming average 500 words per page)
    estimated_pages = total_words // 400

    if estimated_pages > 30:
        return f"DOCX is too long (estimated {estimated_pages} pages, over 30 page limit)"
    
    text = "\n".join([para.text for para in doc.paragraphs])
    summary = await summarize_text(text)
    embedding = await generate_embedding(summary)
    await store_in_vector_db(embedding, summary, username)
    return f"DOCX processed and stored (estimated {estimated_pages} pages)"

async def process_photo(photo, username):
    file = await photo.get_file()
    summary = await extract_text_from_image(file.file_path)
    embedding = await generate_embedding(summary)
    await store_in_vector_db(embedding, summary, username)
    return f"Photo analyzed and stored: {summary}"

async def process_audio(audio, username):
    try:
        file = await audio.get_file()
        file_content = await file.download_as_bytearray()
        
        # Convert MP3 to WAV
        audio_segment = AudioSegment.from_mp3(BytesIO(file_content))
        wav_io = BytesIO()
        audio_segment.export(wav_io, format="wav")
        wav_io.seek(0)
        
        recognizer = sr.Recognizer()
        with sr.AudioFile(wav_io) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)
        
        summary = await summarize_text(text)
        embedding = await generate_embedding(summary)
        await store_in_vector_db(embedding, summary, username)
        return f"Audio transcribed and stored. Transcription: {text}"
    except sr.UnknownValueError:
        return "Speech Recognition could not understand the audio"
    except sr.RequestError as e:
        return f"Could not request results from Speech Recognition service; {e}"
    except Exception as e:
        return f"Error processing audio: {str(e)}"

async def process_text(text, username):
    if text.lower().startswith(("what", "how", "why", "when", "where", "who")):
        return await query_knowledge_base(text, username)
    embedding = await generate_embedding(text)
    await store_in_vector_db(embedding, text, username)
    return "Text processed and stored"

async def process_url(url, username):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    text = soup.get_text()
    print(text)
    summary = await summarize_text(text)
    embedding = await generate_embedding(f"A link was provided ({url}) and here's the summary of what's in that link:\n{summary}")
    await store_in_vector_db(embedding, summary, username)
    return "URL content extracted and stored"